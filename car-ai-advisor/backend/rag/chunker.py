"""
多格式文档分块器 — 支持 JSON / Markdown / PDF / Word (.docx)
企业级 RAG 知识库预处理模块：
1. 统一解析多种知识库源文件，把原始结构化/非结构化文档切割成检索最小单元 chunk
2. 输出统一 `Document` 数据模型，标准化内容、来源、类型、元数据，无缝供给向量化 embeddings.py
3. 适配业务导购知识库：车辆JSON、用户评价、术语库、操作指南FAQ、行业PDF报告、Word规格说明书
"""

import json
import logging
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
import numpy as np

# 日志：记录解析失败、不支持文件、缺失依赖等告警
logger = logging.getLogger(__name__)


# ============================================================
# 统一数据契约 Document — RAG 全链路标准输出结构
# ============================================================
@dataclass
class Document:
    """
    RAG 管道中的最小检索单元（Chunk 块）
    所有文件解析后，全部统一封装为此对象，消除多格式差异
    """
    # 核心文本：送入Embedding模型做向量、检索时匹配的文本内容
    content: str
    # 来源文件名：追溯这条文本来自哪个原始文件
    source: str
    # 文档类型，区分业务数据，检索时可按类型过滤
    doc_type: str          # vehicle车辆 | review评价 | glossary术语 | guide操作指南 | industry行业文档 | faq问答 | pdf | word
    # 块唯一ID，12位uuid，向量库主键
    chunk_id: str = field(default_factory=lambda: uuid.uuid4().hex[:12])
    # 业务元数据：检索后返回给前端展示过滤条件（品牌、车型、页码、标题、评分等）
    metadata: dict = field(default_factory=dict)
    # 向量数组，由 embeddings.py 批量填充，分块器只负责生成文本不做向量化
    embedding: "Optional[np.ndarray]" = None


# ============================================================
# 1. JSON结构化数据解析：把机器可读JSON转为自然语言文本
# 适用三类结构化知识库：车辆参数、车主评价、专业术语库
# ============================================================
def car_to_text(car: dict) -> tuple[str, dict]:
    """
    将车辆结构化JSON展平为通顺自然语言段落
    返回：(拼接后的检索文本, 过滤用元数据字典)
    作用：大模型只能理解自然语言，纯JSON对象无法做语义检索
    """
    # 分层提取JSON各模块数据
    pt = car.get("powertrain", {})       # 动力
    perf = car.get("performance", {})     # 性能
    dims = car.get("dimensions", {})     # 尺寸座位
    sd = car.get("smart_driving", {})     # 智能驾驶
    interior = car.get("interior", {})    # 内饰配置

    parts = [
        f"【{car['full_name']}】",
        f"品牌：{car['brand']}，车型：{car['model']}，类别：{car['category']}",
        f"价格区间：{car['price_range']}，目标用户：{car.get('target_users', '暂无')}",
    ]

    # 动力系统段落：纯电/插混参数统一拼接
    if pt.get("type"):
        powertrain_lines = [f"动力类型：{pt['type']}"]
        if pt.get("battery_type"):
            powertrain_lines.append(f"电池类型：{pt['battery_type']}")
        if pt.get("battery_capacity_kwh"):
            powertrain_lines.append(f"电池容量：{pt['battery_capacity_kwh']}kWh")
        if pt.get("cltc_range_km"):
            powertrain_lines.append(f"CLTC续航：{pt['cltc_range_km']}km")
        if pt.get("pure_electric_range_km"):
            powertrain_lines.append(f"纯电续航：{pt['pure_electric_range_km']}km")
        if pt.get("combined_range_km"):
            powertrain_lines.append(f"综合续航：{pt['combined_range_km']}km")
        if pt.get("fuel_consumption_l_per_100km"):
            powertrain_lines.append(f"馈电油耗：{pt['fuel_consumption_l_per_100km']}L/100km")
        if pt.get("fast_charge"):
            powertrain_lines.append(f"快充：{pt['fast_charge']}")
        if pt.get("energy_density_wh_per_kg"):
            powertrain_lines.append(f"电池能量密度：{pt['energy_density_wh_per_kg']}Wh/kg")
        parts.append("，".join(powertrain_lines))

    # 性能参数段落
    if perf:
        perf_parts = []
        if perf.get("motor_power_kw"):
            perf_parts.append(f"电机功率：{perf['motor_power_kw']}kW")
        if perf.get("system_power_kw"):
            perf_parts.append(f"系统功率：{perf['system_power_kw']}kW")
        if perf.get("zero_to_hundred_seconds"):
            perf_parts.append(f"0-100km/h：{perf['zero_to_hundred_seconds']}s")
        if perf.get("drive_type"):
            perf_parts.append(f"驱动方式：{perf['drive_type']}")
        parts.append("，".join(perf_parts))

    # 车身尺寸段落
    if dims:
        dims_parts = []
        if dims.get("length_mm"):
            dims_parts.append(f"长{dims['length_mm']}mm")
        if dims.get("width_mm"):
            dims_parts.append(f"宽{dims['width_mm']}mm")
        if dims.get("height_mm"):
            dims_parts.append(f"高{dims['height_mm']}mm")
        if dims.get("wheelbase_mm"):
            dims_parts.append(f"轴距{dims['wheelbase_mm']}mm")
        if dims.get("seats"):
            dims_parts.append(f"{dims['seats']}座")
        if dims.get("trunk_volume_l"):
            dims_parts.append(f"后备箱{dims['trunk_volume_l']}L")
        parts.append("尺寸：" + "，".join(dims_parts))

    # 智能驾驶配置
    if sd:
        sd_parts = []
        if sd.get("level"):
            sd_parts.append(f"智驾等级：{sd['level']}")
        if sd.get("chip"):
            sd_parts.append(f"芯片：{sd['chip']}")
        if sd.get("lidar") is True or sd.get("lidar_count"):
            count = sd.get("lidar_count", 1)
            sd_parts.append(f"激光雷达：{count}个")
        if sd.get("highway_noa"):
            sd_parts.append("支持高速NOA")
        if sd.get("city_noa"):
            sd_parts.append("支持城市NOA")
        if sd.get("auto_parking"):
            sd_parts.append("支持自动泊车")
        parts.append("智能驾驶：" + "，".join(sd_parts))

    # 内饰舒适配置
    if interior:
        interior_parts = []
        if interior.get("screen_size_inch"):
            interior_parts.append(f"中控屏{interior['screen_size_inch']}英寸")
        if interior.get("hud"):
            interior_parts.append("HUD抬头显示")
        if interior.get("sound_system"):
            interior_parts.append(interior["sound_system"])
        if interior.get("seat_material"):
            interior_parts.append(interior["seat_material"])
        if interior.get("seat_heating"):
            interior_parts.append("座椅加热")
        if interior.get("seat_ventilation"):
            interior_parts.append("座椅通风")
        if interior.get("ambient_light"):
            interior_parts.append("氛围灯")
        parts.append("内饰：" + "，".join(interior_parts))

    # 附加卖点、质保、竞品
    if car.get("key_features"):
        parts.append("核心亮点：" + "、".join(car["key_features"]))
    if car.get("warranty"):
        parts.append(f"质保：{car['warranty']}")
    if car.get("competitors"):
        parts.append(f"竞品：{'、'.join(car['competitors'])}")

    # 拼接完整自然语言段落
    text = "。".join(parts) + "。"

    # 构建检索过滤元数据，检索后前端可筛选品牌/车型
    metadata = {
        "brand": car.get("brand"),
        "model": car.get("model"),
        "full_name": car.get("full_name"),
        "category": car.get("category"),
        "price_range": car.get("price_range"),
        "powertrain_type": pt.get("type"),
    }
    return text, metadata


def review_to_text(review: dict) -> tuple[str, dict]:
    """车主评价JSON转可读文本，用于用户真实口碑检索。"""
    owner = review.get("owner", "")
    if isinstance(owner, dict):
        owner_text = f"{owner.get('ownership_duration', '未知')}、{owner.get('mileage', '未知')}"
    else:
        owner_text = str(owner) if owner else "未知"
    text = (
        f"车主评价：{review['model']}｜"
        f"评分：{review['rating']}/5.0｜"
        f"用车：{owner_text}｜"
        f"优点：{'、'.join(review.get('pros', []))}｜"
        f"缺点：{'、'.join(review.get('cons', []))}｜"
        f"建议：{review.get('advice', '暂无')}"
    )
    metadata = {"model": review.get("model"), "rating": review.get("rating"), "date": review.get("date")}
    return text, metadata


def glossary_to_text(term: dict) -> tuple[str, dict]:
    """专业术语字典转解释文本，用于知识库专业名词问答"""
    text = f"{term['term']}（{term['full_name']}）：{term['explanation']}"
    metadata = {"term": term["term"], "category": term.get("category")}
    return text, metadata


def chunk_json(file_path: str) -> list[Document]:
    """
    JSON文件统一入口：自动识别JSON业务类型（车辆/评价/术语）
    自动遍历数组每条数据，每条生成一个独立Document块
    """
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    source = Path(file_path).name  # 保存文件名来源
    documents: list[Document] = []

    # 数组类型：车辆列表 / 评价列表
    if isinstance(data, list) and data:
        sample = data[0]
        if "brand" in sample:
            # 车辆数据集
            for item in data:
                text, meta = car_to_text(item)
                documents.append(Document(content=text, source=source, doc_type="vehicle", metadata=meta))
        elif "pros" in sample:
            # 车主评价数据集
            for item in data:
                text, meta = review_to_text(item)
                documents.append(Document(content=text, source=source, doc_type="review", metadata=meta))
    # 字典包含terms字段：术语库
    elif isinstance(data, dict) and "terms" in data:
        for item in data["terms"]:
            text, meta = glossary_to_text(item)
            documents.append(Document(content=text, source=source, doc_type="glossary", metadata=meta))
    else:
        logger.warning(f"无法识别 JSON 类型，跳过: {file_path}")

    return documents


# ============================================================
# 2. Markdown 文件分块：按标题层级切割章节
# 规则：FAQ用三级标题###分块，指南/行业文档用二级标题##分块
# ============================================================
def chunk_markdown(file_path: str) -> list[Document]:
    """
    非结构化MD文档切割：
    1. 根据文件名自动区分文档业务类型 faq/guide/industry
    2. 正则前瞻分割标题，每个标题下的完整章节为独立Chunk
    3. 元数据记录章节标题、序号，检索后可展示来源章节
    """
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    source = Path(file_path).name
    path_lower = file_path.lower()

    # 根据文件名判断文档类型与分割标题层级
    if "faq" in path_lower:
        heading_level = "###"
        doc_type = "faq"
    elif "guide" in path_lower:
        heading_level = "##"
        doc_type = "guide"
    elif "industry" in path_lower:
        heading_level = "##"
        doc_type = "industry"
    else:
        heading_level = "##"
        doc_type = "markdown"

    # 正则前瞻零宽断言：在每个标题前分割，保留标题内容
    pattern = rf"(?=^{heading_level} )"
    sections = re.split(pattern, text, flags=re.MULTILINE)

    documents: list[Document] = []
    for i, section in enumerate(sections):
        section = section.strip()
        if not section:
            continue
        # 清除md分割线---
        section = re.sub(r"^---\s*\n", "", section, flags=re.MULTILINE)
        # 提取章节标题存入metadata
        heading_match = re.match(rf"^{heading_level}\s+(.+)$", section, re.MULTILINE)
        heading = heading_match.group(1).strip() if heading_match else None
        # 文件开头无标题内容标记为前言
        if not heading and i == 0 and not section.startswith("#"):
            heading = "(前言)"

        documents.append(Document(
            content=section,
            source=source,
            doc_type=doc_type,
            metadata={"heading": heading, "section_index": i},
        ))

    # 极端情况：全文无任何标题，整份文档作为单个chunk
    if not documents and text.strip():
        documents.append(Document(
            content=text.strip(), source=source, doc_type=doc_type,
            metadata={"heading": None, "section_index": 0},
        ))

    return documents


# ============================================================
# 3. PDF 文件解析分块（企业报告、产品手册）
# 依赖可选包 pdfplumber，无依赖直接跳过不阻断流程
# ============================================================
def chunk_pdf(file_path: str) -> list[Document]:
    """
    PDF解析逻辑：
    1. 逐页提取文字，标记页码
    2. 优先按中文一、二、三 / 数字1. 标题分块
    3. 无标题则按双换行切割，限制单块最大1000字符，避免超长向量
    4. 扫描件图片PDF无文字则直接返回空列表
    """
    # 懒导入依赖，未安装不抛出崩溃异常，仅打日志
    try:
        import pdfplumber
    except ImportError:
        logger.warning(f"pdfplumber 未安装，跳过 PDF 文件: {file_path}。安装: pip install pdfplumber")
        return []

    source = Path(file_path).name
    full_text_parts: list[str] = []

    try:
        with pdfplumber.open(file_path) as pdf:
            # 逐页提取文本，带上页码标记
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    full_text_parts.append(f"[第{page_num}页]\n{page_text.strip()}")
    except Exception as e:
        logger.error(f"PDF 解析失败: {file_path} — {e}")
        return []

    # 全文无文字（扫描件）
    if not full_text_parts:
        logger.warning(f"PDF 可能为扫描件/图片型，文本提取为空: {file_path}")
        return []

    full_text = "\n\n".join(full_text_parts)

    # 分层分割策略：先中文序号 → 阿拉伯序号 → 固定长度截断
    sections = re.split(r"(?=\n[一二三四五六七八九十]+、)", full_text)
    if len(sections) <= 1:
        sections = re.split(r"(?=\n\d+\.\s*\S)", full_text)
    if len(sections) <= 1:
        # 按段落合并，控制单块长度不超过1000字符
        paragraphs = full_text.split("\n\n")
        chunks: list[str] = []
        current = ""
        for para in paragraphs:
            if len(current) + len(para) < 1000:
                current += "\n\n" + para if current else para
            else:
                if current:
                    chunks.append(current)
                current = para
        if current:
            chunks.append(current)
        sections = chunks

    documents: list[Document] = []
    for i, section in enumerate(sections):
        section = section.strip()
        if not section:
            continue
        # 提取块内页码存入元数据，检索结果展示来源页码
        page_match = re.search(r"\[第(\d+)页\]", section)
        page_hint = int(page_match.group(1)) if page_match else None
        documents.append(Document(
            content=section,
            source=source,
            doc_type="pdf",
            metadata={"section": i, "page": page_hint},
        ))

    return documents


# ============================================================
# 4. Word docx 文档解析（产品规格书、培训文档）
# 依赖 python-docx，支持识别Word自带标题样式、提取表格内容
# ============================================================
def chunk_docx(file_path: str) -> list[Document]:
    """
    Word分块逻辑：
    1. 识别Word内置Heading标题样式，按标题自动分章节
    2. 单独提取表格文本，拼接进文档内容
    3. 无标题则全文合并为单个块
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning(f"python-docx 未安装，跳过 Word 文件: {file_path}。安装: pip install python-docx")
        return []

    source = Path(file_path).name
    sections: list[dict] = []  # 存储分块章节：标题 + 正文行

    try:
        doc = DocxDocument(file_path)

        # 遍历所有段落，区分标题/普通正文
        current = {"heading": None, "body_parts": []}
        for para in doc.paragraphs:
            style = (para.style.name or "").lower()
            text = para.text.strip()
            if not text:
                continue

            # 检测标题样式，开启新分块
            if "heading" in style or "heading" in (para.style.name or "").lower():
                if current["body_parts"] or current["heading"]:
                    sections.append(current)
                # 提取标题层级数字
                heading_level = 1
                for word in style.split():
                    if word.isdigit():
                        heading_level = int(word)
                        break
                current = {"heading": text, "body_parts": [], "heading_level": heading_level}
            else:
                current["body_parts"].append(text)

        # 存入最后一块
        if current["body_parts"] or current["heading"]:
            sections.append(current)

        # 提取Word内所有表格，转为文本
        tables_text: list[str] = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            tables_text.append("\n".join(rows))

        documents: list[Document] = []
        # 组装分块内容
        for i, sec in enumerate(sections):
            body = "\n".join(sec["body_parts"])
            if not body and not sec["heading"]:
                continue
            content = f"## {sec['heading']}\n\n{body}" if sec["heading"] else body
            documents.append(Document(
                content=content.strip(),
                source=source,
                doc_type="word",
                metadata={"heading": sec.get("heading"), "section_index": i},
            ))

        # 无标题章节时，全文合并+表格生成单个块
        if not documents:
            all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if tables_text:
                all_text += "\n\n---\n表格数据：\n" + "\n\n".join(tables_text)
            if all_text.strip():
                documents.append(Document(
                    content=all_text.strip(), source=source, doc_type="word",
                    metadata={"heading": None},
                ))

        return documents

    except Exception as e:
        logger.error(f"DOCX 解析失败: {file_path} — {e}")
        return []


# ============================================================
# 统一对外调度入口
# ============================================================
# 扩展名与对应解析函数映射表，方便后续新增文件类型
SUPPORTED_EXTENSIONS = {
    ".json": chunk_json,
    ".md": chunk_markdown,
    ".pdf": chunk_pdf,
    ".docx": chunk_docx,
}


def chunk_file(file_path: str) -> list[Document]:
    """
    对外统一入口函数，上层预处理管道只调用这一个方法
    自动根据文件后缀匹配解析函数，返回标准化Document列表
    """
    ext = Path(file_path).suffix.lower()
    handler = SUPPORTED_EXTENSIONS.get(ext)
    if handler is None:
        logger.warning(f"不支持的文件格式: {ext}，跳过 {file_path}")
        return []
    # 调用对应分块器
    return handler(file_path)